# backend/app/main.py
import io
import os
import logging
import re
from typing import Optional, List
from copy import deepcopy
import zipfile
import base64

from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware

import mammoth
from markdownify import markdownify as mdify
import pdfplumber
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer

# constants
MAX_BYTES = 3 * 1024 * 1024  # 3 MB
ALLOWED_EXT = {".docx", ".pdf"}

app = FastAPI(title="MD Utility")

# Allow requests from your Next.js frontend (adjust origin in production)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Next.js default dev port
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static (favicon etc)
app.mount(
    "/static",
    StaticFiles(directory=os.path.join(os.path.dirname(__file__), "static")),
    name="static",
)

logger = logging.getLogger("uvicorn.error")


@app.get("/", response_class=HTMLResponse)
async def root():
    html = """
    <html>
      <head><title>MD Utility API</title></head>
      <body>
        <h2>Welcome — FastAPI is running</h2>
        <p>Try the API docs: <a href="/docs">OpenAPI docs</a></p>
      </body>
    </html>
    """
    return HTMLResponse(content=html)


@app.get("/static/favicon.ico")
async def favicon():
    favicon_path = os.path.join(os.path.dirname(__file__), "static", "favicon.ico")
    if os.path.exists(favicon_path):
        return FileResponse(favicon_path)
    return JSONResponse({"detail": "favicon not found"}, status_code=404)


from fastapi import Response

@app.get("/api/download/sample-theme")
def download_sample_theme():
    """
    Download the sample theme file reliably (Sync read).
    """
    path = os.path.join(os.path.dirname(__file__), "static", "sample_theme.docx")
    if not os.path.exists(path):
         raise HTTPException(status_code=404, detail="Sample theme not found")
    
    with open(path, "rb") as f:
        content = f.read()

    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        headers={"Content-Disposition": "attachment; filename=sample_theme.docx"}
    )


@app.get("/health")
async def health():
    return {"status": "ok", "python_version": "3.14"}


# Simple upload echo (kept)
@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    contents = await file.read()
    return {"filename": file.filename, "size_bytes": len(contents)}


# --- Conversion helpers ---
def _file_extension(filename: str) -> str:
    if not filename or "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[1].lower()


async def _read_upload_file(upload_file: UploadFile) -> bytes:
    contents = await upload_file.read()
    return contents


def convert_docx_bytes_to_markdown(contents: bytes) -> str:
    """
    Convert docx bytes -> HTML via mammoth -> Markdown via markdownify.
    """
    with io.BytesIO(contents) as f:
        result = mammoth.convert_to_html(f)
        html = result.value or ""
    # Convert HTML to markdown; adjust options if you want different styles
    md = mdify(html, heading_style="ATX")
    # Optional simple cleanup
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md.strip()


def convert_pdf_bytes_to_markdown(contents: bytes) -> str:
    """
    Extract text from PDF using pdfplumber and join pages.
    Best-effort; complex PDFs may need specialized handling.
    """
    text_parts = []
    with pdfplumber.open(io.BytesIO(contents)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text()
            if txt:
                text_parts.append(txt.strip())
    joined = "\n\n".join(text_parts).strip()
    joined = re.sub(r"\n{3,}", "\n\n", joined)
    return joined


# --- API endpoint for conversion ---
@app.post("/api/convert/word-to-md")
async def convert_word_to_md(file: UploadFile = File(...)):
    filename = file.filename or ""
    ext = _file_extension(filename)
    if ext not in ALLOWED_EXT:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}. Allowed: {', '.join(sorted(ALLOWED_EXT))}")

    contents = await _read_upload_file(file)
    size = len(contents)
    if size == 0:
        raise HTTPException(status_code=400, detail="Empty file uploaded")
    if size > MAX_BYTES:
        raise HTTPException(status_code=413, detail="File too large. Max allowed size is 3 MB.")

    try:
        if ext == ".docx":
            markdown = convert_docx_bytes_to_markdown(contents)
        elif ext == ".pdf":
            markdown = convert_pdf_bytes_to_markdown(contents)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        if markdown is None:
            markdown = ""

        return JSONResponse({"markdown": markdown})
    except Exception as e:
        logger.exception("Conversion failed")
        # Return a sanitized message (avoid leaking stack traces)
        raise HTTPException(status_code=500, detail="Conversion failed: " + str(e))


# --- MD to Word Conversion ---

def add_toc(doc):
    """
    Insert a Table of Contents (TOC) at the start of the document.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # 1. Begin Field
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    fldChar.set(qn('w:dirty'), 'true')
    run._r.append(fldChar)
    
    # 2. Field Instructions
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    # 3. Separate Field
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    # 4. End Field
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    # Page break after TOC (Page 2 done, now Page 3 starts)
    doc.add_page_break()

def force_update_fields(doc):
    """
    Inject w:updateFields into settings.xml to force Word to update TOC/fields on open.
    """
    settings = doc.settings.element
    # Check if element already exists
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    
    update_fields.set(qn('w:val'), 'true')

def force_print_layout(doc):
    """
    Inject w:view w:val="print" into settings.xml to force Print Layout view.
    """
    settings = doc.settings.element
    view = settings.find(qn('w:view'))
    if view is None:
        view = OxmlElement('w:view')
        settings.append(view)
    view.set(qn('w:val'), 'print')   


def _add_html_to_document(doc, html_content: str):
    """
    Basic HTML to Docx parser.
    Supports headings, paragraphs, and lists.
    Forcibly applies 'Segoe UI' to maintain consistency.
    """
    soup = BeautifulSoup(html_content, "html.parser")
    
    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol", "pre", "blockquote", "table"]):
        if element.name.startswith("h"):
            level = int(element.name[1])
            h = doc.add_heading(element.get_text(), level=level)
            # Force Segoe UI
            for run in h.runs:
                run.font.name = 'Segoe UI'
                
        elif element.name == "p":
            p = doc.add_paragraph(element.get_text())
            for run in p.runs:
                run.font.name = 'Segoe UI'
                
        elif element.name in ["ul", "ol"]:
            for li in element.find_all("li"):
                style = "List Bullet" if element.name == "ul" else "List Number"
                p = doc.add_paragraph(li.get_text(), style=style)
                for run in p.runs:
                    run.font.name = 'Segoe UI'
                    
        elif element.name == "pre":
             # Code blocks
             p = doc.add_paragraph()
             runner = p.add_run(element.get_text())
             runner.font.name = 'Courier New' # Code stays Courier
             runner.font.size = Pt(10)
             
        elif element.name == "blockquote":
            p = doc.add_paragraph(element.get_text())
            p.style = 'Quote' # or 'Intense Quote' if available, otherwise fallback
            for run in p.runs:
                run.font.name = 'Segoe UI'
        elif element.name == "table":
            # Basic table handling
            rows = element.find_all("tr")
            if not rows:
                continue
            
            # Determine max columns
            max_cols = 0
            for r in rows:
                cols = r.find_all(["td", "th"])
                max_cols = max(max_cols, len(cols))
            
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'
            
            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                
                # Check if this is a header row (first row or <th> elements)
                is_header = (i == 0)
                
                # Apply row formatting if header
                if is_header:
                    tr = table.rows[i]._tr
                    trPr = tr.get_or_add_trPr()
                    # Cannot easily set row height/color here on row level for all cells in python-docx 
                    # efficiently without iterating cells, so handled in cell loop
                
                for j, cell in enumerate(cells):
                    if j < max_cols:
                        docx_cell = table.cell(i, j)
                        docx_cell.text = cell.get_text().strip()
                        
                        # Custom Formatting
                        # Header Row: BG #003366, Text White, Bold, Center
                        if is_header:
                            # SHADING
                            tcPr = docx_cell._tc.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), '003366') # Dark Blue
                            tcPr.append(shd)
                            
                            # PARAGRAPH / RUN
                            for paragraph in docx_cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.bold = False # Reset then set
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                    run.font.name = 'Segoe UI'
                                    run.font.size = Pt(11)
                        else:
                            # Body Cells: Segoe UI 10.5
                            for paragraph in docx_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Segoe UI'
                                    run.font.size = Pt(10.5)

def _clear_document_content(doc: Document):
    """
    Remove all content from the document body, 
    but preserve the last paragraph to maintain valid XML and Section Properties.
    """
    # Delete all tables
    for tbl in doc.tables:
        tbl._element.getparent().remove(tbl._element)
        
    # Delete all paragraphs EXCEPT the last one
    if len(doc.paragraphs) > 0:
        for p in doc.paragraphs[:-1]:
            p._element.getparent().remove(p._element)
            
        # Clear the last paragraph instead of deleting it
        last_p = doc.paragraphs[-1]
        last_p.clear() 
    else:
        # Should not happen in a valid docx, but safety
        pass

def get_default_document():
    """
    Load the default 'sample_theme.docx' and clear its content
    to serve as the base template with Headers/Footers.
    """
    try:
        theme_path = os.path.join(os.path.dirname(__file__), "static", "sample_theme.docx")
        if os.path.exists(theme_path):
            doc = Document(theme_path)
            _clear_document_content(doc)
            return doc
    except Exception as e:
        print(f"Error loading default theme: {e}")
    
    return Document() # Fallback to blank if missing

def convert_md_to_docx_bytes(md_content: str, theme_bytes: Optional[bytes] = None, theme_is_docx: bool = False) -> bytes:
    # 1. Prepare Main Body Document
    body_doc = None
    
    if theme_is_docx and theme_bytes:
        try:
            body_doc = Document(io.BytesIO(theme_bytes))
            _clear_document_content(body_doc)
        except Exception:
            pass
            
    if body_doc is None:
        body_doc = get_default_document()
        if theme_bytes and not theme_is_docx:
            try:
                theme_text = theme_bytes.decode('utf-8', errors='ignore')
                body_doc.add_paragraph(theme_text)
                body_doc.add_page_break()
            except:
                pass

    # 1a. Prepare Body Doc Structure (Buffer Section + Content Section)
    # Critical: Do this BEFORE adding content so content goes into the (new) S2.
    
    # Check if we have sections (loaded from theme)
    if len(body_doc.sections) > 0:
        s1 = body_doc.sections[0]
        s2 = body_doc.add_section(WD_SECTION.NEW_PAGE)
        s2.header.is_linked_to_previous = False
        s2.footer.is_linked_to_previous = False
        
        # Helper to copy
        def copy_hf(src, tgt):
            if tgt._element.getchildren():
                for child in list(tgt._element):
                    tgt._element.remove(child)
            for child in src._element.getchildren():
                 tgt._element.append(deepcopy(child))

        copy_hf(s1.header, s2.header)
        copy_hf(s1.footer, s2.footer)
    
    # Convert MD to HTML
    html = markdown.markdown(md_content, extensions=['tables'])
    
    # Add TOC to Body Doc (Appends to S2)
    add_toc(body_doc)
    
    # Force update fields on open
    force_update_fields(body_doc)
    
    # Force Print Layout
    force_print_layout(body_doc)
    
    # Add content to Body Doc (Appends to S2)
    _add_html_to_document(body_doc, html)
    
    # 2. Load Cover Page
    cover_path = os.path.join(os.path.dirname(__file__), "static", "Cover_page.docx")
    print(f"Looking for Cover Page at: {cover_path}")
    if os.path.exists(cover_path):
        print("Found Cover Page.")
        master_doc = Document(cover_path)
    else:
        # Fallback if cover missing: Start with Body
        master_doc = body_doc
        body_doc = None # To avoid re-appending if we just made it master

    # 3. Compose Documents
    if body_doc:
        # Pre-process Body Doc to prevent Header Merge
        # We add a buffer section at the start. 
        # S1 (Buffer) -> Merges with Cover (Blank Header)
        # S2 (Content) -> Appended as New Section (Preserves Theme Header)
        composer = Composer(master_doc)
        composer.append(body_doc)
    else:
        print("No Cover Page found, using Body as Master.")
        # If no cover, master is body
        composer = Composer(master_doc)

    # 4. Load & Append End Page
    end_path = os.path.join(os.path.dirname(__file__), "static", "end_page.docx")
    print(f"Looking for End Page at: {end_path}")
    if os.path.exists(end_path):
        try:
            print("Found End Page. Appending...")
            end_doc = Document(end_path)
            # Remove any initial empty paragraphs from end_doc to avoid extra whitespace if needed
            # But usually we just append.
            # Force new section logic is handled by composer if end_doc has 1 section?
            # We want End Page to be a NEW section.
            # If docxcompose merges...
            # We can use the same Buffer strategy? 
            # OR, since master_doc now ends with S2 (Body), and S2 has headers.
            # If End Doc merge into S2 -> It is part of Body. Headers OK.
            # But we want End Doc to be separate section?
            # User said: "end_page... default header and footer will be added... style will be ignore".
            # If we merge, we keep Body headers.
            # If we add section, we Link to Body.
            # Safest: Add section break to master before appending End?
            # Or use Buffer.
            
            # Let's try appending as is. If it merges into S2, it gets S2 headers (Body headers).
            # This satisfies "End page... header and footer will be added".
            # But "Style will be ignore".
            # If `end_doc` has its own styles, `docxcompose` keeps them usually.
            
            # Revert manual add_section for End Page too
            # master_doc.add_section(WD_SECTION.NEW_PAGE) <-- REMOVED
            composer.append(end_doc)
        except Exception:
            pass

    # 5. Fix Section Headers/Footers
    # Goal: 
    #   Section 0 (Cover): As is
    #   Section 1 (Buffer - merged): Gone/Properties merged into Cover
    #   Section 2 (Body - S2 of body_doc): Unlinked, Custom Headers.
    #   Section 3 (End - merged into S2?? Or New S3?)
    
    # If Body had 2 sections (S1, S2).
    # Master had 1 (Cover).
    # Append Body -> Master has S1(Cover+BodyS1), S2(BodyS2).
    # Append End -> If End has 1 section -> Merges into S2?
    # Then End is part of S2.
    # S2 has Body Headers.
    # So End has Body Headers.
    # This matches requirement!
    
    print(f"Final Document Sections: {len(master_doc.sections)}")
    # We don't need to manually link/unlink post-merge because we did it in preparation.

    # Save to buffer
    out = io.BytesIO()
    composer.save(out)
    return out.getvalue()


@app.post("/api/convert/md-to-word")
async def convert_md_to_word(
    files: List[UploadFile] = File(...),
    theme: Optional[UploadFile] = File(None)
):
    if not files:
         raise HTTPException(status_code=400, detail="No files uploaded")

    # Load theme if present
    theme_bytes = None
    theme_is_docx = False
    if theme:
        ext = _file_extension(theme.filename)
        if ext != ".docx":
            raise HTTPException(status_code=400, detail="Theme file must be a .docx document.")
        theme_bytes = await theme.read()
        theme_is_docx = True

    converted_files = []
    
    # Convert each file
    for f in files:
        content = await f.read()
        # Decode MD
        try:
            md_text = content.decode("utf-8")
        except:
            # Try latin-1 fallback?
            md_text = content.decode("latin-1", errors="ignore")
            
        docx_bytes = convert_md_to_docx_bytes(md_text, theme_bytes, theme_is_docx)
        
        # Output filename
        base_name = os.path.splitext(f.filename)[0]
        out_name = f"{base_name}.docx"
        b64 = base64.b64encode(docx_bytes).decode("utf-8")
        
        converted_files.append({
            "filename": out_name,
            "b64": b64
        })

    # If multiple files, also create a zip
    zip_b64 = None
    if len(converted_files) > 0:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for item in converted_files:
                data = base64.b64decode(item["b64"])
                zf.writestr(item["filename"], data)
        zip_b64 = base64.b64encode(zip_buffer.getvalue()).decode("utf-8")

    return JSONResponse({
        "files": converted_files,
        "zip": zip_b64
    })


# --- Git-less Auto-Update Endpoints (Zip Download) ---
import subprocess
import httpx
import shutil
import tempfile
import json

VERSION_FILE = os.path.join(os.path.dirname(__file__), "version.json")

@app.get("/api/update/check")
async def check_update():
    """
    Checks if the local version (sha) differs from GitHub main.
    """
    try:
        # Fetch remote SHA
        url = "https://api.github.com/repos/arnabseth-dev/mdutility/commits/main"
        async with httpx.AsyncClient() as client:
            resp = await client.get(url)
            resp.raise_for_status()
            remote_data = resp.json()
            remote_sha = remote_data["sha"]
        
        # Get local SHA
        local_sha = ""
        if os.path.exists(VERSION_FILE):
            try:
                with open(VERSION_FILE, "r") as f:
                    data = json.load(f)
                    local_sha = data.get("sha", "")
            except:
                pass # corrupted file, treat as empty

        # If local_sha is empty, we act as if we are on an unknown version.
        # But if we just installed, we might not have the file.
        # Let's say update is available if they differ.
        update_available = (local_sha != remote_sha)
        
        return {
            "update_available": update_available, 
            "local_sha": local_sha, 
            "remote_sha": remote_sha
        }
    except Exception as e:
        logger.exception("Update check error")
        return JSONResponse({"update_available": False, "error": str(e)}, status_code=500)


@app.post("/api/update/execute")
async def execute_update():
    """
    Download zip from GitHub, extract, and overwrite local files.
    Then run install scripts and trigger restart.
    """
    try:
        # 1. Download Zip
        zip_url = "https://github.com/arnabseth-dev/mdutility/archive/refs/heads/main.zip"
        logger.info(f"Downloading update from {zip_url}...")
        
        async with httpx.AsyncClient() as client:
            resp = await client.get(zip_url, follow_redirects=True)
            resp.raise_for_status()
            zip_content = resp.content
            
        # 2. Extract and Overwrite
        with tempfile.TemporaryDirectory() as temp_dir:
            zip_path = os.path.join(temp_dir, "update.zip")
            with open(zip_path, "wb") as f:
                f.write(zip_content)
                
            logger.info("Extracting update...")
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # GitHub zips usually have a root folder like 'mdutility-main'
            extract_root = os.path.join(temp_dir, "mdutility-main")
            if not os.path.exists(extract_root):
                # Fallback: find the single directory inside temp_dir
                items = os.listdir(temp_dir)
                items = [i for i in items if os.path.isdir(os.path.join(temp_dir, i))]
                if len(items) == 1:
                    extract_root = os.path.join(temp_dir, items[0])
                else:
                    raise Exception("Unexpected zip structure: could not find root folder")

            # Paths
            # Current working directory is likely 'backend/' (where main.py runs)
            # But let's be safe. We want to identify the project root.
            # We assume __file__ is backend/app/main.py
            # So backend root is backend/
            # Project root is ../
            
            backend_app_dir = os.path.dirname(__file__) # backend/app
            backend_dir = os.path.abspath(os.path.join(backend_app_dir, "..")) # backend
            project_dir = os.path.abspath(os.path.join(backend_dir, "..")) # project root
            
            new_backend = os.path.join(extract_root, "backend")
            new_frontend = os.path.join(extract_root, "frontend")
            
            # --- Detect Changes for Dependencies ---
            has_requirements_change = False
            has_package_json_change = False
            
            def file_content(path):
                if os.path.exists(path):
                    with open(path, "rb") as f: return f.read()
                return None

            # Check requirements.txt (in backend/)
            old_req = file_content(os.path.join(backend_dir, "requirements.txt"))
            new_req = file_content(os.path.join(new_backend, "requirements.txt"))
            if old_req != new_req:
                has_requirements_change = True
            
            # Check package.json (in frontend/)
            frontend_dir = os.path.join(project_dir, "frontend")
            old_pkg = file_content(os.path.join(frontend_dir, "package.json"))
            new_pkg = file_content(os.path.join(new_frontend, "package.json"))
            if old_pkg != new_pkg:
                has_package_json_change = True
            
            # --- Perform Overwrite ---
            logger.info("Overwriting files...")
            
            # Function to recursively copy and overwrite
            def recursive_overwrite(src, dest):
                if not os.path.exists(dest):
                    os.makedirs(dest)
                for item in os.listdir(src):
                    s = os.path.join(src, item)
                    d = os.path.join(dest, item)
                    if os.path.isdir(s):
                        recursive_overwrite(s, d)
                    else:
                        shutil.copy2(s, d)

            # Overwrite Backend
            if os.path.exists(new_backend):
                recursive_overwrite(new_backend, backend_dir)
            
            # Overwrite Frontend
            if os.path.exists(new_frontend):
                recursive_overwrite(new_frontend, frontend_dir)
                
            # --- Update version.json ---
            try:
                sha_url = "https://api.github.com/repos/arnabseth-dev/mdutility/commits/main"
                async with httpx.AsyncClient() as client:
                    r = await client.get(sha_url)
                    if r.status_code == 200:
                        current_sha = r.json()["sha"]
                        with open(VERSION_FILE, "w") as f:
                            json.dump({"sha": current_sha}, f)
            except:
                logger.warning("Failed to update version.json after download")

        # 3. Post-Update Actions (Installs & Rebuilds)
        install_logs = []
        needs_restart = True # Assume code changed, so restart backend
        
        if has_requirements_change:
            logger.info("requirements.txt changed. Running pip install...")
            try:
                subprocess.run(
                    ["pip", "install", "-r", "requirements.txt"],
                    cwd=backend_dir, # Run in backend dir
                    check=True, capture_output=True, text=True
                )
                install_logs.append("Pip install success")
            except subprocess.CalledProcessError as e:
                logger.error(f"Pip install failed: {e.stderr}")
                install_logs.append(f"Pip install failed: {e.stderr}")

        rebuild_msg = ""
        if has_package_json_change:
            logger.info("package.json changed. Running npm install & build...")
            try:
                # npm install
                npm_cmd = "npm.cmd" if os.name == "nt" else "npm"
                subprocess.run([npm_cmd, "install"], cwd=frontend_dir, check=True, shell=True)
                
                # npm run build
                subprocess.run([npm_cmd, "run", "build"], cwd=frontend_dir, check=True, shell=True)
                rebuild_msg = " Frontend rebuilt."
                
                # Restart Frontend Server (Kill 3000)
                try:
                    # Find PID on 3000
                    netstat = subprocess.run(["netstat", "-ano"], capture_output=True, text=True, shell=True)
                    pid_to_kill = None
                    for line in netstat.stdout.splitlines():
                        if ":3000 " in line and "LISTENING" in line:
                            pid_to_kill = line.split()[-1]
                            break
                    
                    if pid_to_kill and pid_to_kill.isdigit() and pid_to_kill != "0":
                         subprocess.run(["taskkill", "/F", "/PID", pid_to_kill], shell=True)
                    
                    # Start new
                    CREATE_NEW_CONSOLE = 0x10
                    subprocess.Popen(
                        [npm_cmd, "start", "--", "-p", "3000"],
                        cwd=frontend_dir,
                        shell=True,
                        creationflags=CREATE_NEW_CONSOLE
                    )
                    rebuild_msg += " Frontend server restarted."
                except Exception as e:
                    rebuild_msg += f" Frontend restart issue: {e}"

                install_logs.append("NPM update success")
            except subprocess.CalledProcessError as e:
                logger.error(f"NPM update failed: {e}")
                install_logs.append(f"NPM update failed: {e}")

        return {
            "success": True, 
            "message": "Update downloaded and applied. " + "; ".join(install_logs) + rebuild_msg, 
            "restart_required": needs_restart,
            "changed_files": ["(Overwritten from Zip)"]
        }

    except Exception as e:
        logger.exception("Update execution error")
        return JSONResponse({"success": False, "message": str(e)}, status_code=500)
